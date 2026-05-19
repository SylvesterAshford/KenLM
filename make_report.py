"""Generate report.docx — a concept + code walkthrough aimed at junior devs."""

from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# ---- base styles ----
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)


def h1(text):
    p = doc.add_heading(text, level=1)
    return p


def h2(text):
    return doc.add_heading(text, level=2)


def h3(text):
    return doc.add_heading(text, level=3)


def para(text, bold=False, italic=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    return p


def code(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = "Consolas"
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0x2B, 0x2B, 0x2B)
    return p


def bullet(text):
    return doc.add_paragraph(text, style="List Bullet")


def numbered(text):
    return doc.add_paragraph(text, style="List Number")


def link_line(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.italic = True
    r.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    return p


# =====================================================================
# TITLE
# =====================================================================
title = doc.add_heading("Domain Adaptation with a Myanmar Language Model", level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("Concept & Code Report — Junior Developer Edition")
r.italic = True
r.font.size = Pt(13)

doc.add_paragraph()  # spacer

para(
    "This report walks through the concepts and code used in class-15.ipynb. "
    "It is written for a junior developer who knows Python but is new to "
    "language modelling. Each concept is introduced in plain English, then "
    "connected to the concept it leads to next, then mapped onto the exact "
    "code that implements it in the notebook."
)

# =====================================================================
# LEARNING MAP
# =====================================================================
h1("1. The learning map — how the concepts connect")

para(
    "Before diving in, here is the chain of ideas. Each concept builds on the "
    "one before it. If a concept feels unclear, the one above it is the place "
    "to revisit."
)

code(
    "  Raw text\n"
    "      ↓ (cleaning)\n"
    "  Clean text\n"
    "      ↓ (sylbreak tokenization)\n"
    "  Stream of syllable tokens\n"
    "      ↓ (count n-grams + Kneser–Ney smoothing)\n"
    "  Language Model (LM)\n"
    "      ↓ (apply to test text)\n"
    "  Per-token probability\n"
    "      ↓ (geometric mean, inverted)\n"
    "  Perplexity (PPL)            ←  our evaluation metric\n"
    "      ↓ (compare across domains)\n"
    "  Domain mismatch identified\n"
    "      ↓ (train tiny in-domain LM + linear interpolation)\n"
    "  Adapted LM with lower PPL"
)

para(
    "Keep this diagram in mind as you read on. Every section of the notebook "
    "implements exactly one step of this pipeline."
)

# =====================================================================
# CONCEPTS
# =====================================================================
h1("2. The concepts you must understand")

# ---- 2.1 LM ----
h2("2.1  What is a language model?")
para(
    "A language model is a function that, given some context, tells you how "
    "likely each possible next word (or syllable) is. Formally, for a sequence "
    "of tokens w₁ … wₙ, an n-gram model assumes each token depends only on "
    "the previous n−1 tokens:"
)
code("P(w₁ … wₙ) ≈ Π P(wᵢ | wᵢ₋ₙ₊₁ … wᵢ₋₁)")
para(
    "We use a 3-gram (trigram) model: each syllable is predicted from the "
    "two syllables before it. KenLM is a fast C++ toolkit that trains and "
    "scores n-gram LMs.",
)
link_line("→ Leads to: we need to count n-grams in text. To count them, we first need tokens. To get tokens, we need clean text.")

# ---- 2.2 cleaning ----
h2("2.2  Why we clean the text first")
para(
    "Real Myanmar text has punctuation, Latin letters, digits, and "
    "occasionally HTML residue. None of that is a Myanmar syllable, and "
    "feeding it into the model would create noise tokens. So before "
    "tokenizing, we strip:"
)
bullet("Myanmar punctuation ။ (ပုဒ်မ) and ၊ (ပုဒ်ထီး)")
bullet("Latin letters and digits (a-z, 0-9)")
bullet("ASCII / Unicode punctuation (! ? , . — “ ” …)")
bullet("Multiple spaces collapse into one")
link_line("→ Leads to: now we have clean Myanmar text. But Myanmar writes with no spaces between words. How do we split it into tokens?")

# ---- 2.3 sylbreak ----
h2("2.3  Sylbreak — tokenizing Myanmar by syllable")
para(
    "English splits on whitespace. Myanmar does not have whitespace between "
    "words, but it has a very regular syllable structure. Ye Kyaw Thu's "
    "sylbreak regex segments Myanmar by syllable: a new syllable starts at "
    "every Myanmar consonant, EXCEPT:"
)
bullet("If the consonant is preceded by virama  ်််  (   ္  )")
bullet("If the consonant itself is followed by virama or asat (်)")
para(
    "Syllables are a clean, language-appropriate unit for Myanmar LMs. "
    "Whether you pick words or syllables matters: perplexity numbers are "
    "only comparable when the tokenizer is identical."
)
link_line("→ Leads to: now we have a stream of syllable tokens. We can finally count n-grams and build a language model.")

# ---- 2.4 smoothing ----
h2("2.4  Smoothing — why raw counts aren't enough")
para(
    "If you only use observed n-gram counts, any unseen n-gram gets "
    "probability zero. That breaks the whole model — a single unseen "
    "trigram in test text makes the whole sentence's probability zero. "
    "Smoothing redistributes some probability mass from observed n-grams to "
    "unseen ones."
)
para(
    "Modified Kneser–Ney smoothing (Chen & Goodman, 1998) is the standard "
    "for n-gram LMs. KenLM's lmplz uses it by default. You don't need to "
    "understand the math to use it, but you should know it exists and that "
    "every probability you read out of the model has been smoothed."
)
para("The flag --discount_fallback is only needed for tiny corpora (like ours). On real data, drop it.")
link_line("→ Leads to: we now have a working LM. How do we measure if it's good?")

# ---- 2.5 perplexity ----
h2("2.5  Perplexity — the evaluation metric")
para(
    "Perplexity (PPL) is the LM's average per-token uncertainty, exponentiated. "
    "Intuitively, a PPL of 50 means the model is 'about as confused as if it "
    "had 50 equally-likely choices per token.'"
)
code("PPL(text) = 2^( -1/N · Σ log₂ P(wᵢ | context) )")
para("Two practical rules:")
bullet("Lower PPL = better fit.")
bullet("PPLs are only comparable between models that use the SAME tokenizer and SAME test-token count. That's why our test sets are all exactly 10 sentences × 20 syllables = 200 tokens.")
link_line("→ Leads to: we measure PPL on three different domains. We discover the model is great in-domain but bad out-of-domain.")

# ---- 2.6 domain mismatch ----
h2("2.6  Domain mismatch — why one LM doesn't fit all text")
para(
    "A language model is a frozen snapshot of the vocabulary and patterns "
    "in its training data. The further a test text drifts from that, the "
    "higher the PPL. For Myanmar:"
)
bullet("Religious gāthā uses Pali loanwords (သီလ, မေတ္တာ, နိဗ္ဗာန, ဈာန) absent in general prose.")
bullet("Legal Myanmar uses formal terms (ပုဒ်မ, တရားရုံး, ပြစ်ဒဏ်) and heavy nominal syntax.")
bullet("Facebook Myanmar uses colloquial particles (ဟေ့, ဪ, နော်), slang and abbreviations.")
link_line("→ Leads to: we need a way to improve the LM on each domain without retraining from scratch.")

# ---- 2.7 adaptation ----
h2("2.7  Domain adaptation by linear interpolation")
para(
    "We keep the big base LM, train a tiny in-domain LM on a small amount of "
    "domain text, and blend their probabilities token-by-token:"
)
code("P_mix(w | h) = λ · P_dom(w | h) + (1 − λ) · P_base(w | h)")
para(
    "λ is a number between 0 and 1 that controls trust: λ=0 means trust the "
    "base LM only, λ=1 means trust the tiny domain LM only. The sweet spot "
    "(often around 0.4–0.6) blends both. We sweep λ ∈ {0.1, …, 0.7} and pick "
    "the value that minimises PPL."
)
para(
    "Linear interpolation is the classical baseline for adaptation. It is "
    "fast, easy to implement, and surprisingly effective even with very "
    "little in-domain text."
)
link_line("→ End of concept chain. From here we are just running and measuring.")

# =====================================================================
# CODE WALKTHROUGH
# =====================================================================
h1("3. Code walkthrough — section by section")

para(
    "The notebook is divided into eight numbered sections. Each section "
    "implements exactly one step of the pipeline described in Section 1 "
    "of this report."
)

# ---- 3.0 setup ----
h2("3.0  Setup")
para(
    "Imports the libraries, sets up the paths, reads KENLM_BIN from the "
    "environment, and asserts that the lmplz binary is reachable."
)
code(
    'KENLM_BIN = os.environ.get("KENLM_BIN", "/home/claude/kenlm/build/bin")\n'
    'assert (Path(KENLM_BIN)/"lmplz").exists(), f"lmplz not found at {KENLM_BIN}"'
)
para(
    "If this cell fails on your machine, you either have not built the "
    "KenLM binaries or have not set KENLM_BIN before launching Jupyter. "
    "See the project README for the setup steps."
)

# ---- 3.1 corpus inspection ----
h2("3.1  Section 1 — Corpus inspection")
para(
    "Reads data/train/general_corpus.txt, prints character/line counts, "
    "and shows a sample. Nothing tricky here — just a sanity check that "
    "the input file is present and not empty."
)

# ---- 3.2 cleaning + sylbreak ----
h2("3.2  Section 2 — Cleaning and tokenization")
para("Two helper functions implement the cleaning + sylbreak concepts:")
code(
    'def clean(text: str) -> str:\n'
    '    text = re.sub(r"[A-Za-z0-9]+", " ", text)\n'
    '    text = re.sub(rf"[{re.escape(PUNCT_TO_STRIP)}]", " ", text)\n'
    '    return re.sub(r"\\s+", " ", text).strip()\n'
    '\n'
    'def sylbreak(text: str, sep: str = " ") -> str:\n'
    '    return break_pattern.sub(rf"{sep}\\1", text).strip()'
)
para(
    "tokenize_file streams a file line-by-line, runs clean then sylbreak on "
    "each line, and writes a syllable-tokenized version. Run on the training "
    "corpus, it produces general_corpus.syl.txt."
)
para(
    "Why stream line-by-line instead of loading the whole file? On real "
    "multi-GB corpora, loading everything into RAM is wasteful. The function "
    "is already written to scale up."
)

# ---- 3.3 train base LM ----
h2("3.3  Section 3 — Train the base LM")
para(
    "We call the lmplz binary as a subprocess: stdin is the tokenized corpus, "
    "stdout is the ARPA model file."
)
code(
    'cmd = [f"{KENLM_BIN}/lmplz", "-o", "3", "--discount_fallback", "-S", "30%"]\n'
    'with open(train_syl) as fin, open(arpa, "w") as fout:\n'
    '    r = subprocess.run(cmd, stdin=fin, stdout=fout, stderr=subprocess.PIPE)'
)
para("Flag meanings:")
bullet("-o 3 — train a trigram model.")
bullet("-S 30% — use up to 30% of available RAM during sorting/counting.")
bullet("--discount_fallback — required for tiny corpora; drop on real data.")
para(
    "Then build_binary converts the human-readable ARPA into a memory-mapped "
    ".bin for fast loading. Finally kenlm.Model(...) loads the binary into "
    "Python so we can score with it."
)

# ---- 3.4 test sets ----
h2("3.4  Section 4 — Domain test sets")
para(
    "Three raw test files (religious_raw.txt, legal_raw.txt, facebook_raw.txt) "
    "are tokenized with the SAME sylbreak code, then trimmed to exactly "
    "200 syllable tokens each (10 sentences × 20 syllables)."
)
para(
    "Why exactly the same length? Perplexity averages logprob across tokens. "
    "Different test lengths would give numbers that are not directly "
    "comparable."
)

# ---- 3.5 PPL evaluation ----
h2("3.5  Section 5 — Perplexity per domain")
para("One line does the actual measurement:")
code('base_ppls[dom] = base_lm.perplexity(text)')
para(
    "KenLM's perplexity() adds <s> / </s> markers and averages logprob over "
    "tokens. We also score a slice of the training data as a sanity floor — "
    "the in-domain PPL should always be the lowest of the four bars."
)
para(
    "The bar chart is on a log scale because the religious PPL is roughly "
    "10× larger than the legal one. A linear y-axis would squash the smaller "
    "bars into invisibility."
)

# ---- 3.6 brainstorm ----
h2("3.6  Section 6 — Adaptation strategies (brainstorm)")
para("A markdown table of options:")
bullet("(a) Mix domain data into training — cheap, but dilutes the general LM and forces full retraining.")
bullet("(b) Linear interpolation — what we use. Tunable per domain.")
bullet("(c) Vocabulary expansion — add domain syllables to the LM vocabulary.")
bullet("(d) Class-based LM — group rare words into classes.")
bullet("(e) Fine-tune a neural LM — the modern default, but outside KenLM.")

# ---- 3.7 adaptation experiment ----
h2("3.7  Section 7 — Adaptation experiment")
para(
    "For each domain, train_domain_lm holds out the first 200 syllables "
    "(the test set) and trains a tiny 2-gram LM on the remaining "
    "in-domain syllables. interp_ppl then walks through the test text "
    "token by token:"
)
code(
    "for (bp, _, _), (dp, _, _) in zip(b, d):\n"
    "    p = lam * (10 ** dp) + (1 - lam) * (10 ** bp)\n"
    "    total_logp += math.log10(p); total_toks += 1\n"
    "return 10 ** (-total_logp / total_toks)"
)
para("Reading the code:")
bullet("full_scores() returns log10 probabilities — that's why we do 10**dp to get back to linear probability before mixing.")
bullet("After mixing, we go back to log10 to accumulate (avoids underflow).")
bullet("Final PPL is 10 to the power of negative-average-logprob — standard PPL formula adapted to base 10.")
para(
    "The sweep over λ values picks the best mix per domain. Results are "
    "plotted in lambda_sweep.png and before_after.png in the outputs folder."
)

# =====================================================================
# THINGS A JUNIOR DEV SHOULD ACTUALLY TAKE AWAY
# =====================================================================
h1("4. Key takeaways for a junior dev")

para("If you remember only five things from this report:")
numbered("A language model is just a probability distribution over token sequences. Don't be intimidated — the math is plumbing around counting and smoothing.")
numbered("Tokenization is part of the model. Change the tokenizer and you change the language model. PPLs only compare under identical tokenization.")
numbered("Perplexity is your friend, but it is sensitive. Same tokenizer, same length, same vocab — only then can you compare two PPL numbers.")
numbered("Domain mismatch is the single biggest reason LMs underperform in practice. Always test on data that looks like production traffic.")
numbered("You can fix a lot of domain mismatch with very little in-domain data. Linear interpolation with even 100 syllables of domain text cuts PPL by 70–90%. Bigger fancier techniques are not always needed.")

# =====================================================================
# GLOSSARY
# =====================================================================
h1("5. Glossary")

terms = [
    ("n-gram", "A contiguous sequence of n tokens. A 3-gram of syllables is three consecutive syllables."),
    ("ARPA file", "Plain-text format for n-gram LMs. Human-readable but slow to load."),
    ("Binary LM (.bin)", "Memory-mapped compiled form of an ARPA file. Loads instantly, scores fast."),
    ("Smoothing", "Redistribution of probability mass so unseen n-grams have non-zero probability."),
    ("Kneser–Ney", "A specific (and effective) smoothing recipe. The default in lmplz."),
    ("Perplexity (PPL)", "Average per-token uncertainty, exponentiated. Lower is better."),
    ("OOV (out-of-vocabulary)", "A token in the test set that was never seen during training."),
    ("Domain", "A coherent subset of text — religious verse, legal contracts, social posts, etc."),
    ("Linear interpolation", "Adaptation method: P_mix = λ·P_dom + (1−λ)·P_base."),
    ("Sylbreak", "A regex that segments Myanmar text by syllable boundary."),
]
for term, defn in terms:
    p = doc.add_paragraph()
    r = p.add_run(f"{term} — ")
    r.bold = True
    p.add_run(defn)

# =====================================================================
# FURTHER READING
# =====================================================================
h1("6. Where to go next")

bullet("Jurafsky & Martin, Speech and Language Processing, Chapter 3 (n-gram LMs) — the canonical textbook chapter, free online.")
bullet("Chen & Goodman 1998, 'An Empirical Study of Smoothing Techniques for Language Modeling' — the paper behind modified Kneser–Ney.")
bullet("Kenneth Heafield's KenLM paper (2011) and the github at github.com/kpu/kenlm.")
bullet("For modern neural LMs, the natural next step after this notebook is a small Transformer trained on the same Myanmar corpus — compare its PPL to KenLM's.")

# =====================================================================
# SAVE
# =====================================================================
out = Path(__file__).resolve().parent / "report.docx"
doc.save(str(out))
print(f"Wrote {out}")
